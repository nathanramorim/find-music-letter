"""
find_lyrics.py — Busca letras de músicas no letras.mus.br e gera PDF ou DOCX.

Uso:
    python find_lyrics.py musicas.txt              # Gera DOCX (padrão)
    python find_lyrics.py musicas.txt --format pdf # Gera PDF
    python find_lyrics.py musicas.txt --single repertorio.docx

Formato do arquivo .txt:
    Artista - Música
    (uma entrada por linha, linhas em branco e # são ignoradas)
"""

import sys
import os
import argparse
import re
import time
from urllib.parse import quote

from curl_cffi import requests as cffi_requests
from bs4 import BeautifulSoup
from ddgs import DDGS
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from fpdf import FPDF
from unidecode import unidecode


BASE_URL = "https://www.letras.mus.br"
DELAY_SECONDS = 1.5  # pausa entre requisições para não sobrecarregar o site

# Sessão compartilhada com impersonation Chrome (bypass Cloudflare)
_session: cffi_requests.Session | None = None


def get_session() -> cffi_requests.Session:
    global _session
    if _session is None:
        _session = cffi_requests.Session(impersonate="chrome124")
        try:
            _session.get(BASE_URL, timeout=15)
        except Exception:
            pass
    return _session


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def slugify(text: str) -> str:
    """Converte texto para slug URL-friendly (igual ao letras.mus.br)."""
    text = unidecode(text).lower()
    text = re.sub(r"[^a-z0-9\s-]", "", text)
    text = re.sub(r"[\s]+", "-", text.strip())
    text = re.sub(r"-+", "-", text)
    return text


def _extract_letras_url_from_href(href: str) -> tuple[str, str, str] | tuple[None, None, None]:
    """Extrai (artist_name, song_name, full_url) de um href do letras.mus.br.
    Suporta /artista-musicas/ID/ e /artista/musica/
    """
    match = re.search(r"letras\.mus\.br/([^/?#]+)/([^/?#]+)", href)
    if not match:
        return None, None, None
    artist_slug, song_slug = match.group(1), match.group(2)
    full_url = f"{BASE_URL}/{artist_slug}/{song_slug}/"
    # Remove sufixo -musicas do slug do artista para obter nome legível
    clean_artist = re.sub(r"-musicas$", "", artist_slug)
    artist_name = clean_artist.replace("-", " ").title()
    # Se song_slug for numérico, o nome será preenchido pelo título da página
    song_name = song_slug if song_slug.isdigit() else song_slug.replace("-", " ").title()
    return artist_name, song_name, full_url


def _search_via_letras(query: str) -> tuple[str, str, str] | tuple[None, None, None]:
    """Busca diretamente no letras.mus.br."""
    session = get_session()
    search_url = f"{BASE_URL}/busca/?q={quote(query)}"
    try:
        response = session.get(search_url, timeout=15)
    except Exception:
        return None, None, None

    if response.status_code != 200:
        return None, None, None

    soup = BeautifulSoup(response.text, "html.parser")
    result = soup.select_one("ul.list-nav a[href]")
    if not result:
        result = soup.select_one(".g-link")
    if not result:
        return None, None, None

    href = result["href"]
    artist_name, song_name, full_url = _extract_letras_url_from_href(href)
    if not full_url:
        return None, None, None

    link_text = result.get_text(separator=" ", strip=True)
    if " - " in link_text:
        song_part, artist_part = link_text.split(" - ", 1)
        song_name = song_part.strip()
        artist_name = artist_part.strip()

    return artist_name, song_name, full_url


def _search_via_duckduckgo(query: str) -> tuple[str, str, str] | tuple[None, None, None]:
    """Usa ddgs para encontrar a URL no letras.mus.br via DuckDuckGo."""
    # Sem aspas e sem site: para ampliar resultados, depois filtra por domínio
    search_query = f"{query} letras.mus.br"
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(search_query, max_results=10))
    except Exception as exc:
        print(f"  [ERRO] DuckDuckGo: {exc}")
        return None, None, None

    if not results:
        return None, None, None

    for result in results:
        href = result.get("href", "")
        if "letras.mus.br" not in href:
            continue
        artist_name, song_name, full_url = _extract_letras_url_from_href(href)
        if not full_url:
            continue
        title = result.get("title", "")
        if " - " in title:
            parts = title.split(" - ")
            song_name = parts[0].strip()
            artist_name = parts[-1].strip()
        return artist_name, song_name, full_url

    return None, None, None


def search_song(query: str) -> tuple[str, str, str] | tuple[None, None, None]:
    """
    Busca pelo título (e opcionalmente artista) da música.
    Tenta letras.mus.br e cai para DuckDuckGo se 403.
    Retorna (artist, song_title, url) ou (None, None, None).
    """
    result = _search_via_letras(query)
    if result[0]:
        return result

    print(f"  [INFO] Busca direta bloqueada, tentando DuckDuckGo...")
    time.sleep(1)
    result = _search_via_duckduckgo(query)
    if result[0]:
        return result

    print(f"  [NÃO ENCONTRADO] Nenhum resultado para: {query!r}")
    return None, None, None


def fetch_lyrics(artist: str | None, song: str) -> tuple[str, str, str] | tuple[None, None, None]:
    """
    Retorna (artist_real, titulo_real, letra) ou (None, None, None) se não encontrar.
    Sempre busca a URL real via search (o site usa IDs numéricos, não slugs).
    """
    query = f"{artist} {song}" if artist else song
    found_artist, found_song, url = search_song(query)
    if not url:
        return None, None, None
    artist = found_artist
    song = found_song

    # Usa a página de impressão: HTML simples, sem bloqueio anti-bot
    print_url = url.rstrip("/") + "/print.html"

    session = get_session()
    try:
        response = session.get(print_url, timeout=15)
    except Exception as exc:
        print(f"  [ERRO] Falha na requisição para {print_url}: {exc}")
        return None, None, None

    if response.status_code != 200:
        print(f"  [ERRO] HTTP {response.status_code} para {print_url}")
        return None, None, None

    soup = BeautifulSoup(response.text, "html.parser")

    # Estrutura real da página de impressão:
    # div.page > div.page-header > h1 > a  (título)
    #                             > h2 > a  (artista)
    # div.page > div.page-container        (cada bloco de estrofes)
    #   > div  (linha)  |  <br> (separador de estrofe)

    page_div = soup.find("div", class_="page")
    if not page_div:
        print(f"  [NÃO ENCONTRADO] Estrutura não reconhecida em {print_url}")
        return None, None, None

    header = page_div.find("div", class_="page-header")
    if header:
        h1 = header.find("h1")
        h2 = header.find("h2")
        page_title = h1.get_text(strip=True) if h1 else song
        if not artist and h2:
            artist = h2.get_text(strip=True)
    else:
        page_title = song

    containers = page_div.find_all("div", class_="page-container")
    if not containers:
        print(f"  [NÃO ENCONTRADO] Letra não localizada em {print_url}")
        return None, None, None

    # Cada page-container tem divs (linhas) e <br> (separa estrofes)
    stanzas = []
    for container in containers:
        current_stanza = []
        for child in container.children:
            if getattr(child, "name", None) == "br":
                if current_stanza:
                    stanzas.append("\n".join(current_stanza))
                    current_stanza = []
            elif getattr(child, "name", None) == "div":
                text = child.get_text(strip=True)
                if text:
                    current_stanza.append(text)
        if current_stanza:
            stanzas.append("\n".join(current_stanza))

    full_lyrics = "\n\n".join(stanzas)

    if not full_lyrics.strip():
        return None, None, None

    return artist, page_title, full_lyrics


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _write_song(doc: Document, artist: str, song_title: str, lyrics: str):
    """Escreve uma música no objeto Document."""
    # Título da música
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(song_title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 30, 30)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Artista
    artist_para = doc.add_paragraph()
    artist_run = artist_para.add_run(artist)
    artist_run.font.size = Pt(13)
    artist_run.font.italic = True
    artist_run.font.color.rgb = RGBColor(80, 80, 80)
    artist_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linha divisória (parágrafo com borda)
    divider = doc.add_paragraph()
    pPr = divider._element.get_or_add_pPr()
    pBdr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr")
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "single")
    bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz", "12")
    bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space", "1")
    bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color", "C8C8C8")
    pBdr.append(bottom)

    # Letra
    for line in lyrics.splitlines():
        if line.strip():
            p = doc.add_paragraph(line)
            p_format = p.paragraph_format
            p_format.line_spacing = 1.15
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(40, 40, 40)
        else:
            # Linha em branco para separar estrofes
            doc.add_paragraph()

    # Página nova antes da próxima música
    doc.add_page_break()


def generate_docx(artist: str, song_title: str, lyrics: str, output_path: str):
    """Gera um DOCX individual para uma música."""
    doc = Document()
    _write_song(doc, artist, song_title, lyrics)
    doc.save(output_path)


def generate_combined_docx(songs_data: list[tuple[str, str, str]], output_path: str):
    """Gera um único DOCX com todas as músicas, cada uma começando numa nova página."""
    doc = Document()
    for i, (artist, title, lyrics) in enumerate(songs_data):
        _write_song(doc, artist, title, lyrics)
    # Remove o último page break
    if doc.paragraphs and doc.paragraphs[-1].text == "":
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def safe_text(text: str) -> str:
    """Remove caracteres que a fpdf2 não consegue renderizar em Latin-1."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


class LyricsPDF(FPDF):
    def header(self):
        pass  # sem cabeçalho padrão; usamos título customizado

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Página {self.page_no()}", align="C")


def _write_song_pdf(pdf: FPDF, artist: str, song_title: str, lyrics: str, first: bool = False):
    """Escreve uma música no objeto PDF com letra em 2 colunas."""
    if not first:
        pdf.add_page()

    # Título da música (largura total)
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(30, 30, 30)
    pdf.multi_cell(pdf.epw, 10, safe_text(song_title), align="C")
    pdf.ln(2)

    # Artista (largura total)
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "I", 13)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(pdf.epw, 8, safe_text(artist), align="C")
    pdf.ln(8)

    # Linha divisória
    pdf.set_draw_color(200, 200, 200)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(8)

    # --- Letra em 2 colunas ---
    COL_GAP = 6          # espaço entre colunas (mm)
    col_w = (pdf.epw - COL_GAP) / 2
    line_h = 5.5
    gap_h = 3            # altura da linha em branco entre estrofes

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(40, 40, 40)

    # Calcula a altura disponível para as colunas
    top_y = pdf.get_y()
    page_bottom = pdf.h - pdf.b_margin

    col_x = [pdf.l_margin, pdf.l_margin + col_w + COL_GAP]
    current_col = 0
    y = top_y

    lines = lyrics.splitlines()

    for line in lines:
        is_blank = line.strip() == ""
        h = gap_h if is_blank else line_h

        # Verifica se ultrapassa o fim da coluna/página
        if y + h > page_bottom:
            if current_col == 0:
                # Passa para a coluna 2
                current_col = 1
                y = top_y
            else:
                # Nova página, volta à coluna 1
                pdf.add_page()
                top_y = pdf.get_y()
                page_bottom = pdf.h - pdf.b_margin
                current_col = 0
                y = top_y

        if is_blank:
            y += gap_h
        else:
            pdf.set_xy(col_x[current_col], y)
            y_before = pdf.get_y()
            pdf.multi_cell(col_w, line_h, safe_text(line), align="L")
            y_after = pdf.get_y()
            # Usa a altura real que o multi_cell consumiu
            actual_height = y_after - y_before
            y = y_before + max(actual_height, line_h)

    # Garante que o cursor fique depois do conteúdo gerado
    pdf.set_y(page_bottom)


def generate_pdf(artist: str, song_title: str, lyrics: str, output_path: str):
    """Gera um PDF individual para uma música."""
    pdf = LyricsPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    _write_song_pdf(pdf, artist, song_title, lyrics, first=True)
    pdf.output(output_path)


def generate_combined_pdf(songs_data: list[tuple[str, str, str]], output_path: str):
    """Gera um único PDF com todas as músicas, cada uma começando numa nova página."""
    pdf = LyricsPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    for i, (artist, title, lyrics) in enumerate(songs_data):
        pdf.add_page()
        _write_song_pdf(pdf, artist, title, lyrics, first=True)
    pdf.output(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def parse_songs_file(filepath: str) -> list[tuple[str | None, str]]:
    songs = []
    with open(filepath, encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue
            if " - " in line:
                artist, song = line.split(" - ", 1)
                songs.append((artist.strip(), song.strip()))
            else:
                # Apenas título — artista será descoberto via busca
                songs.append((None, line))
    return songs


def main():
    parser = argparse.ArgumentParser(
        description="Busca letras no letras.mus.br e gera documentos Word ou PDF."
    )
    parser.add_argument("songs_file", help="Arquivo .txt com lista de músicas")
    parser.add_argument(
        "--output",
        default="./letras",
        help="Diretório de saída dos documentos (padrão: ./letras)",
    )
    parser.add_argument(
        "--single",
        metavar="NOME",
        help="Gera um único documento com todas as músicas. Ex: --single repertorio.docx",
    )
    parser.add_argument(
        "--format",
        choices=["docx", "pdf"],
        default="docx",
        help="Formato de saída: docx (padrão) ou pdf",
    )
    args = parser.parse_args()

    if not os.path.isfile(args.songs_file):
        print(f"[ERRO] Arquivo não encontrado: {args.songs_file}")
        sys.exit(1)

    os.makedirs(args.output, exist_ok=True)

    songs = parse_songs_file(args.songs_file)
    if not songs:
        print("[AVISO] Nenhuma música válida encontrada no arquivo.")
        sys.exit(0)

    print(f"\n{len(songs)} música(s) encontrada(s). Iniciando busca...\n")
    ok = 0
    fail = 0
    combined: list[tuple[str, str, str]] = []

    for i, (artist, song) in enumerate(songs, 1):
        label = f"{artist} - {song}" if artist else song
        print(f"[{i}/{len(songs)}] {label}")
        real_artist, title, lyrics = fetch_lyrics(artist, song)

        if lyrics:
            ok += 1
            if args.single:
                combined.append((real_artist or artist or "", title, lyrics))
                print(f"  [OK] Adicionado ao {args.format.upper()} combinado")
            else:
                slug_artist = slugify(real_artist or artist or "desconhecido")
                slug_song = slugify(title or song)
                filename = f"{slug_artist}-{slug_song}.{args.format}"
                output_path = os.path.join(args.output, filename)
                
                if args.format == "docx":
                    generate_docx(real_artist or artist or "", title, lyrics, output_path)
                else:
                    generate_pdf(real_artist or artist or "", title, lyrics, output_path)
                
                print(f"  [OK] {args.format.upper()} salvo: {output_path}")
        else:
            fail += 1

        if i < len(songs):
            time.sleep(DELAY_SECONDS)

    if args.single and combined:
        single_path = os.path.join(args.output, args.single)
        # Adiciona extensão se não tiver
        if not single_path.endswith(f".{args.format}"):
            single_path += f".{args.format}"
        
        if args.format == "docx":
            generate_combined_docx(combined, single_path)
        else:
            generate_combined_pdf(combined, single_path)
        
        print(f"\n[OK] {args.format.upper()} combinado salvo: {single_path}")

    print(f"\n✓ Concluído: {ok} música(s) gerada(s), {fail} falha(s).")
    print(f"  Saída: {os.path.abspath(args.output)}")


if __name__ == "__main__":
    main()
